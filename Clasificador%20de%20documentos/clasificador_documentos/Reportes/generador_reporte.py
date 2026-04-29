"""
Genera el reporte de inconsistencias como documento Word (.docx).

Reemplaza el log de texto plano con un documento formateado con:
  - Logo de la empresa en el encabezado
  - Cabecera del reporte (fecha, remitente, asunto, totales)
  - Un bloque por documento con sus inconsistencias coloreadas por severidad
  - Borrador de respuesta al proveedor al final

Colores por severidad:
  🔴 Alta    → texto rojo
  🟡 Media   → texto naranja oscuro (#C55A11)
  🔵 Baja    → texto azul (#2E75B6)
"""

import os
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------------------------------------------------------------------------
# Rutas
# ---------------------------------------------------------------------------
_LOGO_RUTA = Path(os.getenv(
    "LOGO_EMPRESA_RUTA",
    r"C:\Users\aux22.gg\Desktop\PROYECTOS\logoMercasa.jpeg"
))
_DIR_REGISTROS = Path(__file__).resolve().parent.parent / "Registros"

# ---------------------------------------------------------------------------
# Colores
# ---------------------------------------------------------------------------
_COLOR_ALTA  = RGBColor(0xC0, 0x00, 0x00)   # rojo oscuro
_COLOR_MEDIA = RGBColor(0xC5, 0x5A, 0x11)   # naranja oscuro
_COLOR_BAJA  = RGBColor(0x2E, 0x75, 0xB6)   # azul corporativo
_COLOR_GRIS  = RGBColor(0x40, 0x40, 0x40)   # texto secundario
_COLOR_FONDO_HEADER = "1F3864"              # azul muy oscuro (hex sin #)

_COLORES_SEV = {"alta": _COLOR_ALTA, "media": _COLOR_MEDIA, "baja": _COLOR_BAJA}
_ICONOS_SEV  = {"alta": "●", "media": "●", "baja": "●"}
_ETIQ_SEV    = {"alta": "CRÍTICO", "media": "ADVERTENCIA", "baja": "MENOR"}

# ---------------------------------------------------------------------------
# Helpers de formato
# ---------------------------------------------------------------------------

def _set_cell_bg(cell, hex_color: str):
    """Aplica color de fondo a una celda de tabla."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _add_run(para, text: str, bold=False, italic=False, size_pt=10,
             color: RGBColor | None = None):
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size_pt)
    if color:
        run.font.color.rgb = color
    return run


def _hr(doc: Document, color_hex="BFBFBF"):
    """Línea horizontal usando borde inferior del párrafo."""
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color_hex)
    pBdr.append(bot)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    return p


def _set_margins(doc: Document, top=2.0, bottom=2.0, left=2.5, right=2.5):
    for section in doc.sections:
        section.top_margin    = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin   = Cm(left)
        section.right_margin  = Cm(right)

# ---------------------------------------------------------------------------
# Función principal
# ---------------------------------------------------------------------------

def generar_reporte_word(
    remitente: str,
    asunto: str,
    docs_con_inconsistencias: list[dict],
    sugerencia: str,
    numero_po: str = "",
    nombre_proveedor: str = "",
) -> Path:
    """
    Genera el reporte .docx y retorna la ruta del archivo creado.

    El archivo se guarda en registros/reporte_PO_YYYY-MM-DD_HHMMSS.docx
    """
    import re as _re
    _DIR_REGISTROS.mkdir(exist_ok=True)
    po_safe = _re.sub(r'[\\/*?:"<>|]', "", numero_po).strip() if numero_po else ""
    if po_safe:
        # Versionar: OC-XXXXXXXX_v1.docx, OC-XXXXXXXX_v2.docx, ...
        existing = list(_DIR_REGISTROS.glob(f"OC-{po_safe}_v*.docx"))
        nums_v = []
        for f in existing:
            m = _re.search(r'_v(\d+)\.docx$', f.name)
            if m:
                nums_v.append(int(m.group(1)))
        version    = (max(nums_v) + 1) if nums_v else 1
        nombre_doc = f"OC-{po_safe}_v{version}.docx"
    else:
        ts_archivo = datetime.now().strftime("%Y-%m-%d_%H%M%S")
        nombre_doc = f"reporte_{ts_archivo}.docx"
    ruta_out = _DIR_REGISTROS / nombre_doc

    doc = Document()
    _set_margins(doc)

    # — Estilo base de párrafos —
    estilo = doc.styles["Normal"]
    estilo.font.name = "Calibri"
    estilo.font.size = Pt(10)

    # ── 1. Logo + título en tabla de 2 columnas ──────────────────────────
    tbl = doc.add_table(rows=1, cols=2)
    tbl.autofit = False
    tbl.columns[0].width = Cm(5)
    tbl.columns[1].width = Cm(13)

    # Celda logo
    cel_logo = tbl.cell(0, 0)
    _set_cell_bg(cel_logo, "FFFFFF")
    if _LOGO_RUTA.exists():
        p_logo = cel_logo.paragraphs[0]
        p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run_logo = p_logo.add_run()
        run_logo.add_picture(str(_LOGO_RUTA), width=Cm(4))

    # Celda título
    cel_titulo = tbl.cell(0, 1)
    _set_cell_bg(cel_titulo, _COLOR_FONDO_HEADER)
    p_titulo = cel_titulo.paragraphs[0]
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_titulo.paragraph_format.space_before = Pt(8)
    _add_run(p_titulo, "REPORTE DE INCONSISTENCIAS", bold=True, size_pt=14,
             color=RGBColor(0xFF, 0xFF, 0xFF))
    p_sub = cel_titulo.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p_sub, "Clasificador de Documentos", italic=True, size_pt=10,
             color=RGBColor(0xBF, 0xD7, 0xED))

    doc.add_paragraph()  # espacio

    # ── 2. Cabecera (fecha, remitente, asunto, totales) ──────────────────
    total_incs  = sum(len(d["inconsistencias"]) for d in docs_con_inconsistencias)
    total_altas = sum(
        1 for d in docs_con_inconsistencias
        for i in d["inconsistencias"] if i.get("severidad") == "alta"
    )
    ts_largo = datetime.now().strftime("%d/%m/%Y  %H:%M hrs")

    filas_meta = [
        ("Fecha y hora",  ts_largo),
        ("Proveedor",     nombre_proveedor or "(no identificado)"),
        ("Remitente",     remitente),
        ("Asunto",        asunto),
        ("Documentos",    str(len(docs_con_inconsistencias))),
        ("Problemas",     f"{total_incs}"
                          + (f"  ({total_altas} crítico{'s' if total_altas != 1 else ''})"
                             if total_altas else "")),
    ]

    tbl_meta = doc.add_table(rows=len(filas_meta), cols=2)
    tbl_meta.autofit = False
    tbl_meta.columns[0].width = Cm(3.5)
    tbl_meta.columns[1].width = Cm(14.5)
    for fila_idx, (etiq, valor) in enumerate(filas_meta):
        cel_e = tbl_meta.cell(fila_idx, 0)
        cel_v = tbl_meta.cell(fila_idx, 1)
        _set_cell_bg(cel_e, "EBF3FB")
        p_e = cel_e.paragraphs[0]
        p_v = cel_v.paragraphs[0]
        _add_run(p_e, etiq, bold=True, size_pt=9, color=_COLOR_GRIS)
        _add_run(p_v, valor, size_pt=9)

    doc.add_paragraph()

    # ── 3. Bloque por documento ───────────────────────────────────────────
    for idx, doc_data in enumerate(docs_con_inconsistencias, 1):
        tipo   = doc_data["tipo"]
        nombre = doc_data["nombre_archivo"]
        incs   = doc_data["inconsistencias"]

        # Encabezado del documento
        p_doc = doc.add_paragraph()
        p_doc.paragraph_format.space_before = Pt(6)
        p_doc.paragraph_format.space_after  = Pt(0)
        _add_run(p_doc, f"Documento {idx:02d}  ", bold=True, size_pt=11)
        _add_run(p_doc, nombre, bold=True, size_pt=11, color=RGBColor(0x1F, 0x38, 0x64))

        p_tipo = doc.add_paragraph()
        p_tipo.paragraph_format.space_before = Pt(0)
        p_tipo.paragraph_format.space_after  = Pt(4)
        _add_run(p_tipo, f"{tipo}  ·  {len(incs)} problema{'s' if len(incs) != 1 else ''}",
                 italic=True, size_pt=9, color=_COLOR_GRIS)

        # Tabla de inconsistencias
        tbl_inc = doc.add_table(rows=len(incs), cols=3)
        tbl_inc.autofit = False
        tbl_inc.columns[0].width = Cm(0.8)   # icono
        tbl_inc.columns[1].width = Cm(3.2)   # campo
        tbl_inc.columns[2].width = Cm(14.0)  # descripción

        for i_idx, inc in enumerate(incs):
            sev   = inc.get("severidad", "baja").lower()
            campo = inc.get("campo", "")
            desc  = inc.get("descripcion", "")
            color = _COLORES_SEV.get(sev, _COLOR_BAJA)
            etiq  = _ETIQ_SEV.get(sev, sev.upper())

            cel_icono = tbl_inc.cell(i_idx, 0)
            cel_campo = tbl_inc.cell(i_idx, 1)
            cel_desc  = tbl_inc.cell(i_idx, 2)

            bg = "FFE7E7" if sev == "alta" else ("FFFBE6" if sev == "media" else "EBF3FB")
            _set_cell_bg(cel_icono, bg)
            _set_cell_bg(cel_campo, bg)
            _set_cell_bg(cel_desc,  "FFFFFF")

            p_icono = cel_icono.paragraphs[0]
            p_icono.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_run(p_icono, f"{i_idx + 1}", bold=True, size_pt=10, color=color)

            p_campo = cel_campo.paragraphs[0]
            _add_run(p_campo, f"[{etiq}]\n", bold=True, size_pt=8, color=color)
            _add_run(p_campo, campo, size_pt=9)

            p_desc = cel_desc.paragraphs[0]
            _add_run(p_desc, desc, size_pt=9)

        _hr(doc)
        doc.add_paragraph()

    # ── 4. Borrador de respuesta al proveedor ────────────────────────────
    if sugerencia:
        p_sec = doc.add_paragraph()
        p_sec.paragraph_format.space_before = Pt(8)
        _add_run(p_sec, "Borrador de Respuesta al Proveedor", bold=True, size_pt=12,
                 color=RGBColor(0x1F, 0x38, 0x64))
        _hr(doc, "2E75B6")

        for linea in sugerencia.splitlines():
            s = linea.strip()
            # Encabezado de sección: línea corta, en mayúsculas, termina en ":"
            # Ej: "INVOICE:", "PACKING LIST:", "CO:"
            es_encabezado_doc = (
                s.endswith(":")
                and len(s) <= 40
                and s[:-1].strip().isupper()
            )
            p_l = doc.add_paragraph()
            if es_encabezado_doc:
                p_l.paragraph_format.space_before = Pt(6)
                p_l.paragraph_format.space_after  = Pt(1)
                _add_run(p_l, s.rstrip(":"), bold=True, size_pt=9,
                         color=RGBColor(0x1F, 0x38, 0x64))
            else:
                p_l.paragraph_format.space_before = Pt(1)
                p_l.paragraph_format.space_after  = Pt(1)
                _add_run(p_l, linea, size_pt=9)


    # ── 5. Pie ───────────────────────────────────────────────────────────
    doc.add_paragraph()
    _hr(doc, "BFBFBF")
    p_pie = doc.add_paragraph()
    p_pie.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _add_run(p_pie, f"Generado por Clasificador de Documentos · {ts_largo}",
             size_pt=8, color=_COLOR_GRIS, italic=True)

    doc.save(str(ruta_out))
    return ruta_out
